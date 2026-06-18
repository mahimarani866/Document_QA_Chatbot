from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path


@dataclass(frozen=True)
class Document:
    text: str
    metadata: dict[str, str]


def load_document(path: Path) -> Document:
    suffix = path.suffix.lower()

    if suffix in {".txt", ".md"}:
        text = path.read_text(encoding="utf-8", errors="ignore")
    elif suffix == ".pdf":
        text = _load_pdf(path)
    elif suffix == ".docx":
        text = _load_docx(path)
    else:
        raise ValueError(f"Unsupported document type: {path.suffix}")

    return Document(text=_normalize(text), metadata={"source": path.name})


def load_documents(paths: list[Path]) -> list[Document]:
    return [load_document(path) for path in paths]


def _load_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("Install pypdf to read PDF files.") from exc

    reader = PdfReader(str(path))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(pages)


def _load_docx(path: Path) -> str:
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:
        raise RuntimeError("Install python-docx to read DOCX files.") from exc

    doc = DocxDocument(str(path))
    return "\n".join(paragraph.text for paragraph in doc.paragraphs)


def _normalize(text: str) -> str:
    text = text.replace("\ufeff", "")
    return "\n".join(line.strip() for line in text.splitlines() if line.strip())
